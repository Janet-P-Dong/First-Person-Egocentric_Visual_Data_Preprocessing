"""Convert fetched Feishu wiki Markdown JSON into an APA-style DOCX.

This script expects the JSON output from:
    lark-cli docs +fetch --doc <url> --doc-format markdown --detail simple --json
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SOURCE_JSON = Path("outputs/feishu_wiki_dissertation_results_framework.json")
OUTPUT_DOCX = Path("outputs/feishu_dissertation_results_framework_apa7.docx")


def set_cell_text(paragraph, text: str) -> None:
    paragraph.text = ""
    add_inline_runs(paragraph, text)


def add_inline_runs(paragraph, text: str) -> None:
    """Add a small subset of Markdown inline emphasis to a paragraph."""
    text = text.replace("\\***", "***").replace("\\**", "**").replace("\\*", "*")
    pattern = re.compile(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("***"):
            run = paragraph.add_run(token[3:-3])
            run.bold = True
            run.italic = True
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_after = Pt(0)

    for name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5"]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(12)
        style.font.color.rgb = None
        style.paragraph_format.line_spacing = 2
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(0)

    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.italic = False
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.italic = False
    styles["Heading 3"].font.bold = True
    styles["Heading 3"].font.italic = True
    styles["Heading 4"].font.bold = True
    styles["Heading 4"].font.italic = False
    styles["Heading 5"].font.bold = True
    styles["Heading 5"].font.italic = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_title_page(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2
    run = p.add_run(title.strip())
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    doc.add_paragraph()


def add_heading(doc: Document, level: int, text: str) -> None:
    level = min(max(level, 1), 5)
    p = doc.add_paragraph(style=f"Heading {level}")
    set_cell_text(p, clean_inline(text))
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def clean_inline(text: str) -> str:
    text = text.strip()
    text = re.sub(r"^#+\s*", "", text)
    text = text.replace("<br>", " ")
    text = re.sub(r"</?strong>", "", text)
    text = re.sub(r"</?b>", "", text)
    return text.strip()


def add_paragraph(doc: Document, text: str) -> None:
    text = clean_inline(text)
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_after = Pt(0)
    add_inline_runs(p, text)


def add_placeholder(doc: Document, kind: str, label: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(label)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def convert_content(doc: Document, content: str) -> None:
    title_match = re.search(r"<title>(.*?)</title>", content, flags=re.S)
    title = clean_inline(title_match.group(1)) if title_match else "Dissertation Results Framework"
    add_title_page(doc, title)

    content = re.sub(r"<title>.*?</title>", "", content, flags=re.S)
    lines = content.splitlines()
    image_count = 0
    table_count = 0
    in_grid = False

    pending: list[str] = []

    def flush_pending() -> None:
        nonlocal pending
        if pending:
            add_paragraph(doc, " ".join(x.strip() for x in pending if x.strip()))
            pending = []

    for raw in lines:
        line = raw.strip()
        if not line:
            flush_pending()
            continue

        if line.startswith("<grid"):
            flush_pending()
            in_grid = True
            add_placeholder(doc, "grid", "[Embedded Feishu grid begins.]")
            continue
        if line.startswith("</grid"):
            flush_pending()
            in_grid = False
            add_placeholder(doc, "grid", "[Embedded Feishu grid ends.]")
            continue
        if line.startswith("<column") or line.startswith("</column"):
            continue

        sheet = re.search(r"<sheet[^>]*sheet-id=\"([^\"]+)\"[^>]*token=\"([^\"]+)\"[^>]*></sheet>", line)
        if sheet:
            flush_pending()
            table_count += 1
            add_placeholder(
                doc,
                "table",
                f"[Embedded Feishu Sheet Table {table_count}: sheet-id={sheet.group(1)}, token={sheet.group(2)}. Data block preserved as source marker.]",
            )
            continue

        if line.startswith("![]("):
            flush_pending()
            image_count += 1
            add_placeholder(doc, "figure", f"[Embedded Feishu Figure {image_count}: image preserved as source marker.]")
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            flush_pending()
            level = min(len(heading.group(1)), 5)
            add_heading(doc, level, heading.group(2))
            continue

        # Treat simple numbered outline markers as normal text unless they are
        # already embedded in prose; this preserves the Feishu source sequence.
        pending.append(line)

    flush_pending()


def main() -> None:
    if not SOURCE_JSON.exists():
        sys.exit(f"Missing source JSON: {SOURCE_JSON}")

    data = json.loads(SOURCE_JSON.read_text())
    content = data["data"]["document"]["content"]

    doc = Document()
    configure_document(doc)
    convert_content(doc, content)
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
