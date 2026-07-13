from copy import deepcopy
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Inches, Pt


SOURCE = Path("Dissertation Results Framework - APA 7.docx")
OUTPUT = Path("Dissertation Results Framework - APA 7.docx")

FONT = "Times New Roman"
BODY_SIZE = Pt(12)
TABLE_SIZE = Pt(10)


MAJOR_HEADINGS = {
    "Table of Content",
    "List of Tables and Figures",
    "Abstract",
    "Introduction",
    "Literature Review",
    "Study Question, Purpose and Implications",
    "Conclusion",
    "References",
    "Supplementary",
}


def set_run_font(run, size=BODY_SIZE, bold=None, italic=None):
    run.font.name = FONT
    run.font.size = size
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), FONT)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.underline = False


def set_para_base(paragraph):
    try:
        paragraph.style = "Normal"
    except Exception:
        pass
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.left_indent = None
    pf.right_indent = None
    pf.first_line_indent = Inches(0.5)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        set_run_font(run)


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def add_page_number(paragraph, tabbed=False):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if tabbed else WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.first_line_indent = None
    if tabbed:
        paragraph.paragraph_format.tab_stops.clear_all()
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        tab = paragraph.add_run("\t")
        set_run_font(tab)
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_run_font(run)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def paragraph_after(paragraph, text=""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if text:
        new_para.add_run(text)
    return new_para


def set_text(paragraph, text):
    clear_paragraph(paragraph)
    paragraph.add_run(text)


def table_caption_split(text):
    m = re.match(r"^(Table\s+(?:[A-Za-z]|\d+(?:\.\d+)*|x|X))\.\s+(.+)$", text)
    if m:
        return m.group(1), m.group(2)
    m = re.match(r"^(Table\s+(?:[A-Za-z]|\d+(?:\.\d+)*|x|X))\s+(.+)$", text)
    if m and len(m.group(2).split()) > 3:
        return m.group(1), m.group(2)
    return None


def classify(text, index):
    compact = " ".join(text.split())
    if not compact:
        return "blank"
    if index == 0:
        return "title"
    if table_caption_split(compact):
        return "table_combined"
    if re.match(r"^Table\s+[\w.]+$", compact, re.I):
        return "table_number"
    if re.match(r"^Figure\s+.+", compact, re.I):
        return "figure_number"
    if compact.startswith("Note.") or compact.startswith("*Note."):
        return "note"
    if compact in MAJOR_HEADINGS or re.match(r"^Study\s+\d+\b", compact):
        return "level1"
    if re.match(r"^\d+\.\s+\S", compact):
        return "level2"
    if re.match(r"^\d+(?:\.\d+)+\s+\S", compact):
        return "level3"
    if re.match(r"^Supplementary\s+\d", compact):
        return "level3"
    return "body"


def apply_heading(paragraph, level):
    pf = paragraph.paragraph_format
    pf.first_line_indent = None
    pf.left_indent = None
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        set_run_font(run, bold=True, italic=(level == 3))


def apply_caption_number(paragraph):
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        set_run_font(run, bold=True, italic=False)


def apply_caption_title(paragraph):
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        set_run_font(run, bold=False, italic=True)


def apply_note(paragraph):
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        set_run_font(run)
    text = paragraph.text
    if text.startswith("*Note."):
        set_text(paragraph, "Note." + text[len("*Note."):])
        text = paragraph.text
    if text.startswith("Note."):
        clear_paragraph(paragraph)
        r1 = paragraph.add_run("Note.")
        set_run_font(r1, italic=True)
        r2 = paragraph.add_run(text[len("Note."):])
        set_run_font(r2)


def apply_body(paragraph):
    paragraph.paragraph_format.first_line_indent = Inches(0.5)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        set_run_font(run)


def apply_references(paragraph):
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.5)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        set_run_font(run)


def format_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            mar = tc_pr.first_child_found_in("w:tcMar")
            if mar is None:
                mar = OxmlElement("w:tcMar")
                tc_pr.append(mar)
            for side in ("top", "left", "bottom", "right"):
                node = mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    mar.append(node)
                node.set(qn("w:w"), "100")
                node.set(qn("w:type"), "dxa")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = None
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, TABLE_SIZE)


def main():
    doc = Document(SOURCE)

    try:
        normal = doc.styles["Normal"]
    except KeyError:
        normal = None
    if normal is not None:
        normal.font.name = FONT
        normal.font.size = BODY_SIZE
        normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.first_line_indent = Inches(0.5)

    for section in doc.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.different_first_page_header_footer = False
        header = section.header
        for child in list(header._element):
            header._element.remove(child)
        table = header.add_table(rows=1, cols=1, width=Inches(6.5))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        remove_table_borders(table)
        cell = table.cell(0, 0)
        cell.width = Inches(6.5)
        paragraph = cell.paragraphs[0]
        clear_paragraph(paragraph)
        add_page_number(paragraph)

    in_references = False
    paragraphs = list(doc.paragraphs)
    i = 0
    while i < len(paragraphs):
        paragraph = paragraphs[i]
        set_para_base(paragraph)
        text = " ".join(paragraph.text.split())
        kind = classify(text, i)

        if text == "References":
            in_references = True
        elif in_references and (text == "Supplementary" or re.match(r"^Study\s+\d+\b", text)):
            in_references = False

        if kind == "blank":
            paragraph.paragraph_format.first_line_indent = None
        elif kind == "title":
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.space_before = Pt(36)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, bold=True)
        elif kind == "table_combined":
            number, title = table_caption_split(text)
            set_text(paragraph, number)
            apply_caption_number(paragraph)
            new_para = paragraph_after(paragraph, title)
            set_para_base(new_para)
            apply_caption_title(new_para)
            paragraphs.insert(i + 1, new_para)
            i += 1
        elif kind in {"table_number", "figure_number"}:
            apply_caption_number(paragraph)
        elif i > 0 and classify(" ".join(paragraphs[i - 1].text.split()), i - 1) in {"table_number", "figure_number"}:
            apply_caption_title(paragraph)
        elif kind == "note":
            apply_note(paragraph)
        elif kind == "level1":
            apply_heading(paragraph, 1)
        elif kind == "level2":
            apply_heading(paragraph, 2)
        elif kind == "level3":
            apply_heading(paragraph, 3)
        elif in_references:
            apply_references(paragraph)
        else:
            apply_body(paragraph)
        i += 1

    for table in doc.tables:
        format_table(table)

    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
